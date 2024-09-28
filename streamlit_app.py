# import streamlit as st
# import os
# import re
# from docx import Document
# from pdfminer.high_level import extract_text
# from sklearn.feature_extraction.text import TfidfVectorizer
# from sklearn.metrics.pairwise import cosine_similarity

# UPLOAD_FOLDER = "uploads"
# DATABASE_FOLDER = "database"

# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
# os.makedirs(DATABASE_FOLDER, exist_ok=True)

# def extract_text_from_file(file_path):
#     """Extracts text from PDF or DOCX files."""
#     if file_path.endswith(".pdf"):
#         return extract_text(file_path)
#     elif file_path.endswith((".docx", ".doc")):
#         doc = Document(file_path)
#         text = [paragraph.text for paragraph in doc.paragraphs]
#         return "\n".join(text)
#     else:
#         return None

# def simple_sentence_tokenize(text):
#     """A simple sentence tokenizer that splits on periods, question marks, and exclamation points."""
#     return re.split(r'(?<=[.!?])\s+', text)

# def calculate_similarity(text1, text2):
#     """Calculates cosine similarity and highlights similar sentences."""
#     text1_sentences = simple_sentence_tokenize(text1)
#     text2_sentences = simple_sentence_tokenize(text2)

#     # Combine all sentences for vectorization:
#     all_sentences = text1_sentences + text2_sentences

#     # Calculate Similarity using sentences:
#     vectorizer = TfidfVectorizer()
#     tfidf_matrix = vectorizer.fit_transform(all_sentences)

#     # Cosine similarity between all sentences:
#     similarity_matrix = cosine_similarity(tfidf_matrix, tfidf_matrix)

#     # Extract similarity scores for text1 and text2 sentences:
#     similarity_scores = similarity_matrix[
#         : len(text1_sentences), len(text1_sentences) :
#     ]

#     # Highlight Similar Sentences:
#     highlighted_text1 = []
#     highlighted_text2 = []
#     for i, sentence1 in enumerate(text1_sentences):
#         for j, sentence2 in enumerate(text2_sentences):
#             similarity = similarity_scores[i, j]
#             if similarity > 0.8:  # Adjust the threshold if needed
#                 highlighted_text1.append(f"<mark>{sentence1}</mark>")
#                 highlighted_text2.append(f"<mark>{sentence2}</mark>")
#                 break  # Move to the next sentence in text1
#         if len(highlighted_text1) <= i:
#             highlighted_text1.append(sentence1)

#     return (
#         similarity_matrix.max(),
#         " ".join(highlighted_text1),
#         " ".join(highlighted_text2),
#     )

# # --- Streamlit UI ---
# st.title("Plagiarism Detection App")

# # Database Management Section:
# st.header("Manage Database PDFs")

# # Initialize session state for database files:
# if "database_files" not in st.session_state:
#     st.session_state.database_files = os.listdir(DATABASE_FOLDER)

# st.write(
#     f"**Existing Database Files:** {', '.join(st.session_state.database_files) or 'None'}"
# )

# db_file = st.file_uploader("Add a PDF to the database", type=["pdf"])
# if db_file is not None:
#     file_path = os.path.join(DATABASE_FOLDER, db_file.name)
#     with open(file_path, "wb") as f:
#         f.write(db_file.getbuffer())
#     st.success(f"File '{db_file.name}' added to database.")
#     # Update database files in session state:
#     st.session_state.database_files.append(db_file.name)

# # Delete PDF from the database
# file_to_delete = st.selectbox("Select a PDF to delete", st.session_state.database_files)
# if st.button("Delete PDF"):
#     file_path = os.path.join(DATABASE_FOLDER, file_to_delete)
#     if os.path.exists(file_path):
#         os.remove(file_path)
#         st.success(f"File '{file_to_delete}' deleted from database.")
#         # Update database files in session state:
#         st.session_state.database_files.remove(file_to_delete)
#         st.experimental_rerun()  # Refresh Streamlit
#     else:
#         st.error(f"File '{file_to_delete}' not found in the database.")

#     st.markdown("---")  # Visual separator

# # Plagiarism Detection Section:

# st.header("Plagiarism Detection")

# uploaded_file = st.file_uploader(
#     "Upload your document (PDF or DOCX)", type=["pdf", "docx"]
# )

# if uploaded_file is not None:
#     # Save the uploaded file temporarily
#     file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
#     with open(file_path, "wb") as f:
#         f.write(uploaded_file.getbuffer())

#     input_text = extract_text_from_file(file_path)

#     if input_text:
#         similarity_results = []
#         for filename in st.session_state.database_files:
#             database_file_path = os.path.join(DATABASE_FOLDER, filename)
#             database_text = extract_text_from_file(database_file_path)
#             if database_text:
#                 (
#                     similarity,
#                     highlighted_input,
#                     highlighted_database,
#                 ) = calculate_similarity(input_text, database_text)
#                 similarity_results.append(
#                     {
#                         "document": filename,
#                         "similarity": round(similarity * 100, 2),
#                         "highlighted_input": highlighted_input,
#                         "highlighted_database": highlighted_database,
#                     }
#                 )

#         st.subheader("Plagiarism Results:")
#         for result in similarity_results:
#             st.write(f"**Document:** {result['document']}")
#             st.write(f"**Similarity:** {result['similarity']}%")
#             st.markdown(
#                 f"**Highlighted Input Text:** {result['highlighted_input']}",
#                 unsafe_allow_html=True,
#             )
#             st.markdown(
#                 f"**Highlighted Database Text:** {result['highlighted_database']}",
#                 unsafe_allow_html=True,
#             )
#             st.markdown("---")

#     else:
#         st.error("Error: Could not extract text from the uploaded file.")


import streamlit as st
import os
import re
from docx import Document
from pdfminer.high_level import extract_text
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import difflib

UPLOAD_FOLDER = "uploads"
DATABASE_FOLDER = "database"

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DATABASE_FOLDER, exist_ok=True)

def extract_text_from_file(file_path):
    """Extracts text from PDF or DOCX files."""
    if file_path.endswith(".pdf"):
        return extract_text(file_path)
    elif file_path.endswith((".docx", ".doc")):
        doc = Document(file_path)
        text = [paragraph.text for paragraph in doc.paragrapaths]
        return "\n".join(text)
    else:
        return None

def preprocess_text(text):
    """Preprocess the text by removing special characters and extra whitespace."""
    text = re.sub(r'[^\w\s]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.lower().strip()

def calculate_similarity_and_highlight(text1, text2):
    """Calculates cosine similarity and highlights similar text."""
    # Preprocess texts
    processed_text1 = preprocess_text(text1)
    processed_text2 = preprocess_text(text2)

    # Calculate TF-IDF vectors
    vectorizer = TfidfVectorizer()
    tfidf_matrix = vectorizer.fit_transform([processed_text1, processed_text2])

    # Calculate cosine similarity
    similarity = cosine_similarity(tfidf_matrix[0:1], tfidf_matrix[1:2])[0][0]

    # Find and highlight similar parts
    d = difflib.Differ()
    diff = list(d.compare(processed_text1.split(), processed_text2.split()))

    highlighted_text1 = []
    highlighted_text2 = []
    
    for word in diff:
        if word.startswith('  '):  # Common words
            highlighted_text1.append(f'<span style="background-color: yellow;">{word[2:]}</span>')
            highlighted_text2.append(f'<span style="background-color: yellow;">{word[2:]}</span>')
        elif word.startswith('- '):  # Words only in text1
            highlighted_text1.append(word[2:])
        elif word.startswith('+ '):  # Words only in text2
            highlighted_text2.append(word[2:])

    return (
        similarity * 100,  # Convert to percentage
        ' '.join(highlighted_text1),
        ' '.join(highlighted_text2)
    )

# --- Streamlit UI ---
st.title("Plagiarism Detection App")

# Database Management Section:
st.header("Manage Database PDFs")

# Initialize session state for database files:
if "database_files" not in st.session_state:
    st.session_state.database_files = os.listdir(DATABASE_FOLDER)

st.write(
    f"**Existing Database Files:** {', '.join(st.session_state.database_files) or 'None'}"
)

db_file = st.file_uploader("Add a PDF to the database", type=["pdf"])
if db_file is not None:
    file_path = os.path.join(DATABASE_FOLDER, db_file.name)
    with open(file_path, "wb") as f:
        f.write(db_file.getbuffer())
    st.success(f"File '{db_file.name}' added to database.")
    # Update database files in session state:
    st.session_state.database_files.append(db_file.name)

# Delete PDF from the database
file_to_delete = st.selectbox("Select a PDF to delete", st.session_state.database_files)
if st.button("Delete PDF"):
    file_path = os.path.join(DATABASE_FOLDER, file_to_delete)
    if os.path.exists(file_path):
        os.remove(file_path)
        st.success(f"File '{file_to_delete}' deleted from database.")
        # Update database files in session state:
        st.session_state.database_files.remove(file_to_delete)
        st.experimental_rerun()  # Refresh Streamlit
    else:
        st.error(f"File '{file_to_delete}' not found in the database.")

    st.markdown("---")  # Visual separator

# Plagiarism Detection Section:

st.header("Plagiarism Detection")

uploaded_file = st.file_uploader(
    "Upload your document (PDF or DOCX)", type=["pdf", "docx"]
)

if uploaded_file is not None:
    # Save the uploaded file temporarily
    file_path = os.path.join(UPLOAD_FOLDER, uploaded_file.name)
    with open(file_path, "wb") as f:
        f.write(uploaded_file.getbuffer())

    input_text = extract_text_from_file(file_path)

    if input_text:
        similarity_results = []
        for filename in st.session_state.database_files:
            database_file_path = os.path.join(DATABASE_FOLDER, filename)
            database_text = extract_text_from_file(database_file_path)
            if database_text:
                (
                    similarity,
                    highlighted_input,
                    highlighted_database,
                ) = calculate_similarity_and_highlight(input_text, database_text)
                similarity_results.append(
                    {
                        "document": filename,
                        "similarity": round(similarity, 2),
                        "highlighted_input": highlighted_input,
                        "highlighted_database": highlighted_database,
                    }
                )

        st.subheader("Plagiarism Results:")
        for result in similarity_results:
            st.write(f"**Document:** {result['document']}")
            st.write(f"**Similarity:** {result['similarity']}%")
            st.markdown("**Highlighted Input Text:**")
            st.markdown(result['highlighted_input'], unsafe_allow_html=True)
            st.markdown("**Highlighted Database Text:**")
            st.markdown(result['highlighted_database'], unsafe_allow_html=True)
            st.markdown("---")

    else:
        st.error("Error: Could not extract text from the uploaded file.")